"""
report_service.py — Génération de rapports PDF et Word
depuis la session SAMI (profil en mémoire + fitscore)
Utilise : reportlab (PDF) + python-docx (Word)
"""

import io
from datetime import datetime
from typing import Optional

# ── PDF ───────────────────────────────────────────────────────
from reportlab.lib.pagesizes   import A4
from reportlab.lib              import colors
from reportlab.lib.styles       import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units        import cm
from reportlab.platypus         import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak
)
from reportlab.lib.enums        import TA_CENTER, TA_LEFT, TA_RIGHT

# ── Word ──────────────────────────────────────────────────────
from docx                       import Document
from docx.shared                import Pt, RGBColor, Cm, Inches
from docx.enum.text             import WD_ALIGN_PARAGRAPH
from docx.oxml.ns               import qn
from docx.oxml                  import OxmlElement

# ── Couleurs SUPMTI ───────────────────────────────────────────
GREEN  = colors.HexColor("#006666")
RED    = colors.HexColor("#CC0000")
DARK   = colors.HexColor("#1a1a2e")
LIGHT  = colors.HexColor("#f8fafc")
GRAY   = colors.HexColor("#64748b")


# ═══════════════════════════════════════════════════════════════
# HELPER — Extraire données lisibles depuis le profil SAMI
# ═══════════════════════════════════════════════════════════════

def _extraire_donnees(profil: dict, fitscore: dict) -> dict:
    info  = profil.get("informations_personnelles", {}) or {}
    parc  = profil.get("parcours_academique",       {}) or {}
    pref  = profil.get("preferences",               {}) or {}
    psycho= profil.get("profil_psychometrique",     {}) or {}

    prenom   = info.get("prenom", "Étudiant")
    nom      = info.get("nom", "")
    ville    = info.get("ville", "—")
    bac      = parc.get("label_bac") or parc.get("type_bac", "—")
    moyenne  = parc.get("moyenne_generale", 0)
    niveau   = parc.get("niveau_actuel", "—")
    mention  = parc.get("mention", "—")
    interets = pref.get("centres_interet", [])
    ambition = pref.get("ambition_professionnelle", "—")

    classement = []
    if fitscore and fitscore.get("classement"):
        classement = fitscore["classement"][:5]
    meilleure  = fitscore.get("meilleure_filiere", "—") if fitscore else "—"
    rapport_fs = fitscore.get("rapport", "") if fitscore else ""

    points_forts = psycho.get("points_forts", [])

    return {
        "prenom":      prenom,
        "nom":         nom,
        "nom_complet": f"{prenom} {nom}".strip(),
        "ville":       ville,
        "bac":         bac,
        "moyenne":     moyenne,
        "niveau":      niveau,
        "mention":     mention,
        "interets":    interets,
        "ambition":    ambition,
        "classement":  classement,
        "meilleure":   meilleure,
        "rapport_fs":  rapport_fs,
        "points_forts":points_forts,
        "date":        datetime.now().strftime("%d/%m/%Y"),
    }


# ═══════════════════════════════════════════════════════════════
# GÉNÉRATION PDF
# ═══════════════════════════════════════════════════════════════

def generer_rapport_pdf(profil: dict, fitscore: Optional[dict] = None) -> bytes:
    """Retourne le PDF en bytes."""
    d    = _extraire_donnees(profil, fitscore or {})
    buf  = io.BytesIO()
    doc  = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm,  bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()

    # Styles personnalisés
    s_titre = ParagraphStyle("Titre", parent=styles["Title"],
        fontSize=22, textColor=GREEN, spaceAfter=6,
        alignment=TA_CENTER, fontName="Helvetica-Bold")

    s_sous  = ParagraphStyle("Sous", parent=styles["Normal"],
        fontSize=11, textColor=GRAY, spaceAfter=20,
        alignment=TA_CENTER)

    s_h1    = ParagraphStyle("H1", parent=styles["Heading1"],
        fontSize=14, textColor=GREEN, spaceBefore=16, spaceAfter=6,
        fontName="Helvetica-Bold")

    s_h2    = ParagraphStyle("H2", parent=styles["Heading2"],
        fontSize=12, textColor=DARK, spaceBefore=10, spaceAfter=4,
        fontName="Helvetica-Bold")

    s_body  = ParagraphStyle("Body", parent=styles["Normal"],
        fontSize=10, textColor=DARK, spaceAfter=4, leading=15)

    s_label = ParagraphStyle("Label", parent=styles["Normal"],
        fontSize=9, textColor=GRAY, fontName="Helvetica-Oblique")

    story = []

    # ── PAGE DE GARDE ─────────────────────────────────────────
    story.append(Spacer(1, 1.5*cm))
    story.append(Paragraph("SUPMTI Meknès", ParagraphStyle("ecole",
        fontSize=13, textColor=GRAY, alignment=TA_CENTER)))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph("Rapport d'Orientation Académique", s_titre))
    story.append(Paragraph(f"Généré par SAMI IA · {d['date']}", s_sous))
    story.append(HRFlowable(width="100%", thickness=2,
        color=GREEN, spaceAfter=20))

    # ── SECTION 1 : PROFIL ────────────────────────────────────
    story.append(Paragraph("1. Profil Étudiant", s_h1))

    data_profil = [
        ["Nom complet",  d["nom_complet"] or "—",    "Ville",    d["ville"]],
        ["BAC",          d["bac"],                    "Moyenne",  f"{d['moyenne']}/20" if d["moyenne"] else "—"],
        ["Niveau",       d["niveau"],                 "Mention",  d["mention"]],
    ]
    tbl = Table(data_profil, colWidths=[3.5*cm, 5*cm, 3*cm, 5*cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,-1), LIGHT),
        ("BACKGROUND",  (0,0), (0,-1), GREEN),
        ("BACKGROUND",  (2,0), (2,-1), GREEN),
        ("TEXTCOLOR",   (0,0), (0,-1), colors.white),
        ("TEXTCOLOR",   (2,0), (2,-1), colors.white),
        ("FONTNAME",    (0,0), (-1,-1), "Helvetica"),
        ("FONTSIZE",    (0,0), (-1,-1), 9),
        ("PADDING",     (0,0), (-1,-1), 6),
        ("GRID",        (0,0), (-1,-1), 0.5, colors.white),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [LIGHT, colors.HexColor("#e8f5f5")]),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.5*cm))

    if d["interets"]:
        story.append(Paragraph("Centres d'intérêt", s_h2))
        story.append(Paragraph("  ·  ".join(d["interets"]), s_body))

    if d["ambition"] and d["ambition"] != "—":
        story.append(Paragraph("Ambition professionnelle", s_h2))
        story.append(Paragraph(d["ambition"], s_body))

    # ── SECTION 2 : FITSCORE ──────────────────────────────────
    if d["classement"]:
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("2. Classement FitScore", s_h1))
        story.append(Paragraph(
            f"Filière recommandée : <b>{d['meilleure']}</b>", s_body))
        story.append(Spacer(1, 0.3*cm))

        data_fs = [["#", "Filière", "Score de compatibilité"]]
        for i, f in enumerate(d["classement"], 1):
            bar = "█" * int(f["score_total"] / 5)
            data_fs.append([
                str(i),
                f.get("filiere_nom") or f.get("filiere_id", "—"),
                f"{f['score_total']:.1f}%  {bar}"
            ])

        tbl2 = Table(data_fs, colWidths=[1*cm, 7*cm, 8.5*cm])
        tbl2.setStyle(TableStyle([
            ("BACKGROUND",  (0,0), (-1,0),  GREEN),
            ("TEXTCOLOR",   (0,0), (-1,0),  colors.white),
            ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0,0), (-1,-1), 9),
            ("PADDING",     (0,0), (-1,-1), 6),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#e8f5f5")]),
            ("GRID",        (0,0), (-1,-1), 0.3, GRAY),
            ("TEXTCOLOR",   (0,1), (0,-1),  GREEN),
            ("FONTNAME",    (0,1), (0,-1),  "Helvetica-Bold"),
        ]))
        story.append(tbl2)

        if d["rapport_fs"]:
            story.append(Spacer(1, 0.5*cm))
            story.append(Paragraph("Analyse détaillée", s_h2))
            # Limiter le rapport à 1500 chars pour le PDF
            txt = d["rapport_fs"][:1500].replace("\n", "<br/>")
            story.append(Paragraph(txt, s_body))

    # ── SECTION 3 : POINTS FORTS ──────────────────────────────
    if d["points_forts"]:
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("3. Points Forts (Test Psychométrique)", s_h1))
        for pt in d["points_forts"]:
            story.append(Paragraph(f"✓  {pt}", s_body))

    # ── PIED DE PAGE ──────────────────────────────────────────
    story.append(Spacer(1, 1*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=GRAY))
    story.append(Paragraph(
        f"Document généré automatiquement par SAMI · SUPMTI Meknès · {d['date']}",
        ParagraphStyle("footer", fontSize=8, textColor=GRAY, alignment=TA_CENTER)
    ))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════
# GÉNÉRATION WORD
# ═══════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    """Colorier le fond d'une cellule Word."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def generer_rapport_word(profil: dict, fitscore: Optional[dict] = None) -> bytes:
    """Retourne le DOCX en bytes."""
    d   = _extraire_donnees(profil, fitscore or {})
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── EN-TÊTE ───────────────────────────────────────────────
    hdr = doc.add_heading("Rapport d'Orientation Académique", 0)
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hdr.runs:
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x66)
        run.font.size = Pt(22)

    sub = doc.add_paragraph(f"SUPMTI Meknès · SAMI IA · {d['date']}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ── SECTION 1 : PROFIL ────────────────────────────────────
    h1 = doc.add_heading("1. Profil Étudiant", 1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x66)

    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = "Table Grid"
    labels = [
        ("Nom complet", d["nom_complet"] or "—", "Ville",    d["ville"]),
        ("BAC",         d["bac"],                 "Moyenne",  f"{d['moyenne']}/20" if d["moyenne"] else "—"),
        ("Niveau",      d["niveau"],              "Mention",  d["mention"]),
    ]
    for i, (l1, v1, l2, v2) in enumerate(labels):
        row = tbl.rows[i]
        for j, (txt, is_label) in enumerate([(l1,True),(v1,False),(l2,True),(v2,False)]):
            cell = row.cells[j]
            cell.text = txt
            run  = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if is_label:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(cell, "006666")
            else:
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph()

    if d["interets"]:
        doc.add_heading("Centres d'intérêt", 2)
        doc.add_paragraph("  ·  ".join(d["interets"]))

    if d["ambition"] and d["ambition"] != "—":
        doc.add_heading("Ambition professionnelle", 2)
        doc.add_paragraph(d["ambition"])

    # ── SECTION 2 : FITSCORE ──────────────────────────────────
    if d["classement"]:
        doc.add_heading("2. Classement FitScore", 1)
        p = doc.add_paragraph()
        p.add_run("Filière recommandée : ").bold = False
        run = p.add_run(d["meilleure"])
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x66)

        tbl2 = doc.add_table(rows=1 + len(d["classement"]), cols=3)
        tbl2.style = "Table Grid"

        # En-tête
        for j, txt in enumerate(["#", "Filière", "Score"]):
            cell = tbl2.rows[0].cells[j]
            cell.text = txt
            run  = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
            _set_cell_bg(cell, "006666")

        for i, f in enumerate(d["classement"], 1):
            row = tbl2.rows[i]
            row.cells[0].text = str(i)
            row.cells[1].text = f.get("filiere_nom") or f.get("filiere_id", "—")
            row.cells[2].text = f"{f['score_total']:.1f}%"
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)

        if d["rapport_fs"]:
            doc.add_heading("Analyse détaillée", 2)
            doc.add_paragraph(d["rapport_fs"][:1500])

    # ── SECTION 3 : POINTS FORTS ──────────────────────────────
    if d["points_forts"]:
        doc.add_heading("3. Points Forts", 1)
        for pt in d["points_forts"]:
            doc.add_paragraph(pt, style="List Bullet")

    # ── PIED DE PAGE ──────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Document généré par SAMI · SUPMTI Meknès · {d['date']}"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()